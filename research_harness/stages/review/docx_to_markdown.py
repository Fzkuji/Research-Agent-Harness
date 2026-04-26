from __future__ import annotations

from openprogram.agentic_programming.function import agentic_function
from openprogram.agentic_programming.runtime import Runtime


@agentic_function(render_range={"depth": 0, "siblings": 0})
def docx_to_markdown(docx_path: str, runtime: Runtime) -> str:
    """Convert a Word document (.docx, .doc) to clean Markdown for downstream review.

    You have shell access. Try the following extractors in order; fall back
    if the previous is unavailable or produces empty output:

    1. `pandoc -f docx -t gfm --wrap=none` — preferred when pandoc is installed.
       Best heading / list / table fidelity.
       ```bash
       pandoc -f docx -t gfm --wrap=none "$DOCX" -o "$OUT_MD"
       ```

    2. `python-docx` — pure-Python, no external dep beyond pip:
       ```python
       from docx import Document
       doc = Document(docx_path)
       parts = []
       for p in doc.paragraphs:
           style = p.style.name if p.style else ""
           text = p.text
           if not text.strip():
               continue
           if style.startswith("Heading 1"): parts.append(f"# {text}")
           elif style.startswith("Heading 2"): parts.append(f"## {text}")
           elif style.startswith("Heading 3"): parts.append(f"### {text}")
           else: parts.append(text)
       # Tables (basic):
       for tbl in doc.tables:
           for row in tbl.rows:
               parts.append(" | ".join(cell.text for cell in row.cells))
       md = "\\n\\n".join(parts)
       ```

    3. `mammoth` (python lib) as a third fallback:
       ```python
       import mammoth
       with open(docx_path, "rb") as f:
           result = mammoth.convert_to_markdown(f)
           md = result.value
       ```

    4. For legacy `.doc` (binary, not .docx): use `textutil -convert html` (macOS),
       `antiword`, or `libreoffice --headless --convert-to docx` then re-run.

    Post-processing rules (apply to the converted markdown):
    - Collapse runs of 3+ blank lines into 2.
    - Drop the References / Bibliography section entirely (everything after a
      heading matching `(?i)^\\s*#+\\s*(references|bibliography|参考文献)\\s*$`).
    - If pandoc produced raw HTML blocks (figures, comments), strip them or
      replace with `[figure: <caption>]` placeholders.
    - Track-changes / comments markup from Word: keep the accepted text only,
      drop revision marks. With pandoc, use `--track-changes=accept`.

    Sanity checks before returning:
    - Output length must be >= 1500 characters; if smaller, the .docx likely
      failed to extract — try the next extractor.
    - The first 500 characters should contain a plausible title line. If not,
      prepend `# UNKNOWN_TITLE` so downstream stages know extraction was lossy.

    Output file:
    - Save to `<docx_path with .docx/.doc replaced by .md>`. If the same path
      exists, append `.v2`, `.v3`, ... so prior conversions are not overwritten.

    Return value (string):
    `Saved to <md_path>. Extracted <N> chars using <extractor>.`


    # Persistence
    Save your COMPLETE output to a file in the current working directory.
    The filename is dictated by the rules above (mirror the docx_path with .md).
    After saving, return the one-line summary described in "Return value".
    """
    return runtime.exec(content=[
        {"type": "text", "text": f"DOCX path: {docx_path}"},
    ])
